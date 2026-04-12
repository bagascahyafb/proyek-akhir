from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- DOCX GENERATOR ---
def safe_text(value):
    if value is None:
        return ""
    return str(value).strip()

def set_margins(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

def add_section_header(doc, title):
    p = doc.add_paragraph()
    runner = p.add_run(title.upper())
    runner.bold = True
    runner.font.size = Pt(11)
    runner.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_ats_docx(data, language="English"):
    doc = Document()
    set_margins(doc)
    
    # --- FIX 1: SET GLOBAL FONT (Biar gak perlu set ulang tiap paragraf) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    lang_map = {
        "summary": "Professional Summary" if language == "English" else "Ringkasan Profesional",
        "exp": "Work Experience" if language == "English" else "Pengalaman Kerja",
        "edu": "Education" if language == "English" else "Pendidikan",
        "proj": "Projects" if language == "English" else "Proyek",
        "hard_skill": "Hard Skills" if language == "English" else "Keahlian Teknis",
        "soft_skill": "Soft Skills" if language == "English" else "Keahlian Interpersonal",
        "cert": "Certifications" if language == "English" else "Sertifikasi",
        "award": "Honors & Awards" if language == "English" else "Penghargaan",
    }

    # 1. HEADER
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Nama besar dan tebal
    name_run = h1.add_run(data.get('Personal_Info', {}).get('Nama', 'Unknown User'))
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.name = 'Times New Roman'
    
    # Kontak
    contact_list = [x for x in [data['Personal_Info']['HP'], data['Personal_Info']['Email'], data['Personal_Info']['LinkedIn'], data['Personal_Info']['Alamat'], data['Personal_Info']['Portfolio']] if x]
    h2 = doc.add_paragraph(" | ".join(contact_list))
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Hapus manual style font name disini karena udah ikut global, cukup size aja
    h2.runs[0].font.size = Pt(10) 
    
    # 2. SUMMARY
    if data['Personal_Info']['Summary']:
        add_section_header(doc, lang_map['summary'])
        p = doc.add_paragraph(data['Personal_Info']['Summary'])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Gak perlu set font name/italic lagi, otomatis ikut Normal (TNR Regular)

    # 3. EXPERIENCE
    if data['Experience']:
        add_section_header(doc, lang_map['exp'])
        for item in data['Experience']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            # Posisi (Bold)
            run_pos = p.add_run(f"{item['Posisi']}")
            run_pos.bold = True
            # Perusahaan
            p.add_run(f" | {item['Perusahaan']}")
            # --- FIX 2: TANGGAL ITALIC (Pake add_run biar gak ngerusak global) ---
            p_date = doc.add_paragraph()
            p_date.paragraph_format.space_after = Pt(2)
            run_date = p_date.add_run(item['Durasi'])
            run_date.italic = True # Cuma run ini yang miring!
            run_date.font.size = Pt(10)
            
            # Deskripsi
            if 'Deskripsi' in item:
                for line in item['Deskripsi'].split('\n'):
                    if line.strip():
                        # Style 'List Bullet' biasanya aman, tapi kita pastikan font-nya
                        p_desc = doc.add_paragraph(line.strip().replace('- ', ''), style='List Bullet')
                        p_desc.paragraph_format.space_after = Pt(0)
                        if p_desc.runs:
                            p_desc.runs[0].font.name = 'Times New Roman'

    # 4. PROJECTS
    if data['Projects']:
        add_section_header(doc, lang_map['proj'])
        for item in data['Projects']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            
            run_name = p.add_run(f"{item['Nama_Proyek']}")
            run_name.bold = True
            
            meta = []
            if item.get('Role'): meta.append(item['Role'])
            if item.get('Tech_Stack'): meta.append(f"({item['Tech_Stack']})")
            p_date = doc.add_paragraph()
            p_date.paragraph_format.space_after = Pt(2)
            run_date = p_date.add_run(item['Duration'])
            run_date.italic = True 
            run_date.font.size = Pt(10)

            if meta: p.add_run(f" | {' '.join(meta)}")
            
            if 'Deskripsi' in item:
                for line in item['Deskripsi'].split('\n'):
                    if line.strip():
                        p_desc = doc.add_paragraph(line.strip().replace('- ', ''), style='List Bullet')
                        p_desc.paragraph_format.space_after = Pt(0)
                        if p_desc.runs:
                            p_desc.runs[0].font.name = 'Times New Roman'

    # 5. EDUCATION
    if data['Education']:
        add_section_header(doc, lang_map['edu'])
        for item in data['Education']:
            institusi = safe_text(item.get('Institusi'))
            tahun_lulus = safe_text(item.get('Tahun_Lulus'))
            gelar = safe_text(item.get('Gelar'))
            jurusan = safe_text(item.get('Jurusan'))
            ipk = safe_text(item.get('IPK'))

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            
            run_uni = p.add_run(institusi)
            run_uni.bold = True
            
            if tahun_lulus:
                p.add_run(f"  ({tahun_lulus})")
            
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(6)
            
            edu_parts = [part for part in [gelar, jurusan] if part]
            if edu_parts:
                degree_line = " in ".join(edu_parts) if len(edu_parts) == 2 else edu_parts[0]
                run_major = p2.add_run(degree_line)
                run_major.italic = True
            
            if ipk:
                label = "GPA" if language == "English" else "IPK"
                p2.add_run(f" | {label}: {ipk}")

    # 6. SKILLS
    if data['Skills_Hard'] or data['Skills_Soft']:
        add_section_header(doc, "Skills")
        if data['Skills_Hard']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            runner = p.add_run(f"{lang_map['hard_skill']}: ")
            runner.bold = True
            p.add_run(", ".join(data['Skills_Hard']))
        
        if data['Skills_Soft']:
            p = doc.add_paragraph()
            runner = p.add_run(f"{lang_map['soft_skill']}: ")
            runner.bold = True
            p.add_run(", ".join(data['Skills_Soft']))

    # 7. AWARDS
    items = []
    for c in data['Certifications']: items.append(f"{c['Nama']} - {c['Penerbit']} ({c['Tahun']})")
    for a in data['Awards']: items.append(f"Award: {a['Nama_Award']} - {a['Pemberi']} ({a['Tahun']})")
    
    if items:
        add_section_header(doc, f"{lang_map['cert']} & {lang_map['award']}")
        for it in items:
            p = doc.add_paragraph(it, style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
            if p.runs:
                p.runs[0].font.name = 'Times New Roman'

    return doc
