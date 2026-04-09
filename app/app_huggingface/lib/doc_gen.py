from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def safe_text(value):
    if value is None:
        return ""
    return str(value).strip()


def set_margins(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def add_section_header(doc, title):
    paragraph = doc.add_paragraph()
    runner = paragraph.add_run(title.upper())
    runner.bold = True
    runner.font.size = Pt(11)
    runner.font.name = "Times New Roman"
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(4)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def generate_ats_docx(data, language="English"):
    doc = Document()
    set_margins(doc)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    lang_map = {
        "summary": "Professional Summary" if language == "English" else "Ringkasan Profesional",
        "exp": "Work Experience" if language == "English" else "Pengalaman Kerja",
        "edu": "Education" if language == "English" else "Pendidikan",
        "proj": "Projects" if language == "English" else "Proyek",
        "hard_skill": "Hard Skills" if language == "English" else "Keahlian Teknis",
        "soft_skill": "Soft Skills" if language == "English" else "Keahlian Interpersonal",
        "cert": "Certifications" if language == "English" else "Sertifikasi",
        "award": "Honors & Awards" if language == "English" else "Penghargaan",
    }

    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = h1.add_run(data["Personal_Info"]["Nama"])
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.name = "Times New Roman"

    contact_list = [
        value
        for value in [
            data["Personal_Info"]["HP"],
            data["Personal_Info"]["Email"],
            data["Personal_Info"]["LinkedIn"],
            data["Personal_Info"]["Alamat"],
        ]
        if value
    ]
    h2 = doc.add_paragraph(" | ".join(contact_list))
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.runs[0].font.size = Pt(10)

    if data["Personal_Info"]["Summary"]:
        add_section_header(doc, lang_map["summary"])
        paragraph = doc.add_paragraph(data["Personal_Info"]["Summary"])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if data["Experience"]:
        add_section_header(doc, lang_map["exp"])
        for item in data["Experience"]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            run_pos = paragraph.add_run(f"{item['Posisi']}")
            run_pos.bold = True
            paragraph.add_run(f" | {item['Perusahaan']}")

            p_date = doc.add_paragraph()
            p_date.paragraph_format.space_after = Pt(2)
            run_date = p_date.add_run(item["Durasi"])
            run_date.italic = True
            run_date.font.size = Pt(10)

            if "Deskripsi" in item:
                for line in item["Deskripsi"].split("\n"):
                    if line.strip():
                        p_desc = doc.add_paragraph(line.strip().replace("- ", ""), style="List Bullet")
                        p_desc.paragraph_format.space_after = Pt(0)
                        if p_desc.runs:
                            p_desc.runs[0].font.name = "Times New Roman"

    if data["Projects"]:
        add_section_header(doc, lang_map["proj"])
        for item in data["Projects"]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)

            run_name = paragraph.add_run(f"{item['Nama_Proyek']}")
            run_name.bold = True

            meta = []
            if item.get("Role"):
                meta.append(item["Role"])
            if item.get("Tech_Stack"):
                meta.append(f"({item['Tech_Stack']})")

            p_date = doc.add_paragraph()
            p_date.paragraph_format.space_after = Pt(2)
            run_date = p_date.add_run(item["Duration"])
            run_date.italic = True
            run_date.font.size = Pt(10)

            if meta:
                paragraph.add_run(f" | {' '.join(meta)}")

            if "Deskripsi" in item:
                for line in item["Deskripsi"].split("\n"):
                    if line.strip():
                        p_desc = doc.add_paragraph(line.strip().replace("- ", ""), style="List Bullet")
                        p_desc.paragraph_format.space_after = Pt(0)
                        if p_desc.runs:
                            p_desc.runs[0].font.name = "Times New Roman"

    if data["Education"]:
        add_section_header(doc, lang_map["edu"])
        for item in data["Education"]:
            institusi = safe_text(item.get("Institusi"))
            tahun_lulus = safe_text(item.get("Tahun_Lulus"))
            gelar = safe_text(item.get("Gelar"))
            jurusan = safe_text(item.get("Jurusan"))
            ipk = safe_text(item.get("IPK"))

            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)

            run_uni = paragraph.add_run(institusi)
            run_uni.bold = True

            if tahun_lulus:
                paragraph.add_run(f"  ({tahun_lulus})")

            second = doc.add_paragraph()
            second.paragraph_format.space_after = Pt(6)

            edu_parts = [part for part in [gelar, jurusan] if part]
            if edu_parts:
                degree_line = " in ".join(edu_parts) if len(edu_parts) == 2 else edu_parts[0]
                run_major = second.add_run(degree_line)
                run_major.italic = True

            if ipk:
                label = "GPA" if language == "English" else "IPK"
                second.add_run(f" | {label}: {ipk}")

    if data["Skills_Hard"] or data["Skills_Soft"]:
        add_section_header(doc, "Skills")
        if data["Skills_Hard"]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            runner = paragraph.add_run(f"{lang_map['hard_skill']}: ")
            runner.bold = True
            paragraph.add_run(", ".join(data["Skills_Hard"]))

        if data["Skills_Soft"]:
            paragraph = doc.add_paragraph()
            runner = paragraph.add_run(f"{lang_map['soft_skill']}: ")
            runner.bold = True
            paragraph.add_run(", ".join(data["Skills_Soft"]))

    items = []
    for cert in data["Certifications"]:
        items.append(f"{cert['Nama']} - {cert['Penerbit']} ({cert['Tahun']})")
    for award in data["Awards"]:
        items.append(f"Award: {award['Nama_Award']} - {award['Pemberi']} ({award['Tahun']})")

    if items:
        add_section_header(doc, f"{lang_map['cert']} & {lang_map['award']}")
        for item in items:
            paragraph = doc.add_paragraph(item, style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(0)
            if paragraph.runs:
                paragraph.runs[0].font.name = "Times New Roman"

    return doc
